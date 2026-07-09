"""
extractor.py
Extracts raw text from uploaded resume files (PDF or DOCX).
"""

import io
import re
import fitz  # PyMuPDF
import docx  # python-docx
from docx.document import Document as _DocxDocument
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph


def _iter_block_items(parent):
    """
    Yield each paragraph and table child, IN THE ORDER they appear in the
    document body. python-docx's `document.paragraphs` alone SKIPS any
    paragraph that lives inside a table cell - which means resumes using a
    table for a two-column/sidebar layout (a very common modern template
    style) would lose that entire column's content. This walks the raw XML
    body so nothing gets silently dropped.
    """
    parent_elm = parent.element.body if isinstance(parent, _DocxDocument) else parent._tc
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def _normalize_bullets(text: str) -> str:
    """
    Word-generated PDFs/DOCX often use Wingdings/Symbol fonts for bullet
    points. Those bullets don't map to a real Unicode bullet character -
    they land in the Private Use Area (U+E000-U+F8FF) instead, which shows
    up as garbled boxes ('?', '☒', etc.) once extracted and displayed in a
    normal font. Since resumes essentially never use PUA characters for
    anything OTHER than bullets, it's safe to normalize all of them to '•'.
    """
    return re.sub(r"[\uE000-\uF8FF]", "•", text)


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF given as raw bytes."""
    text_parts = []
    with fitz.open(stream=file_bytes, filetype="pdf") as doc:
        for page in doc:
            text_parts.append(page.get_text())
    return _normalize_bullets("\n".join(text_parts))


def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extract text from a DOCX given as raw bytes, including table content
    (e.g. sidebar/two-column resume templates that use a table for layout).
    """
    document = docx.Document(io.BytesIO(file_bytes))
    parts = []

    def _add_paragraph(p):
        if p.text.strip():
            parts.append(p.text.strip())

    def _add_table(table):
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _add_paragraph(p)

    for block in _iter_block_items(document):
        if isinstance(block, Paragraph):
            _add_paragraph(block)
        elif isinstance(block, Table):
            _add_table(block)

    return _normalize_bullets("\n".join(parts))


def extract_text(uploaded_file) -> str:
    """
    Dispatcher: accepts a Streamlit UploadedFile object and routes it
    to the correct extractor based on file extension.
    """
    filename = uploaded_file.name.lower()
    file_bytes = uploaded_file.read()

    if filename.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    elif filename.endswith(".docx"):
        return extract_text_from_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {filename}. Please upload a PDF or DOCX.")
